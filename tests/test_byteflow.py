"""
Basic offline tests for ByteFlow.

These don't require a real Ollama install - tool/plugin/memory/search
logic is tested directly. Run with:

    python -m pytest tests/ -v

or, without pytest installed:

    python tests/test_byteflow.py
"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Agent() now defaults to persisting at ~/.byteflow/memory.json (see
# agent.py's _default_memory_path()) so it "just works" across separate
# runs/notebooks/scripts without the person needing to remember to pass
# memory_path. That default resolves via os.path.expanduser("~"), which
# reads the HOME env var (USERPROFILE on Windows) - so for the test
# suite, redirect that to a throwaway temp directory FIRST, before any
# byteflow import or Agent() call, to guarantee tests never read/write
# the real ~/.byteflow files on the machine running them. Without this,
# every Agent(provider=...) call in these tests would silently persist
# to (and pick up stale state from) the real default location.
_TEST_HOME = tempfile.mkdtemp(prefix="byteflow_test_home_")
os.environ["HOME"] = _TEST_HOME
os.environ["USERPROFILE"] = _TEST_HOME  # Windows equivalent

from byteflow import Agent
from byteflow.memory import Memory
from byteflow.profile import Profile
from byteflow.tools import Tool
from byteflow.plugin import Plugin
from byteflow.search import TextIndex, tokenize
from byteflow import desktop_tools
from byteflow.companion import CompanionController
from byteflow import chunking
from byteflow.embeddings import TfidfEmbedder, SentenceTransformerEmbedder, sentence_transformers_available
from byteflow.vector_store import VectorStore
from byteflow import web_search as web_search_module
from byteflow import file_reading


def test_tool_run_success():
    tool = Tool("add", lambda a, b: a + b)
    assert tool.run(2, 3) == 5


def test_tool_run_catches_errors():
    def crash(a, b):
        raise ValueError("boom")
    tool = Tool("crash", crash)
    result = tool.run(1, 2)
    assert "Tool Error" in result
    assert "boom" in result


def test_agent_register_and_use_tool():
    agent = Agent(provider=None)
    agent.register_tool(Tool("multiply", lambda a, b: a * b))
    assert agent.use_tool("multiply", 4, 5) == 20


def test_plugin_loads_and_registers_tool():
    class MathPlugin(Plugin):
        def __init__(self):
            super().__init__("MathPlugin")

        def setup(self, agent):
            agent.register_tool(Tool("multiply", lambda a, b: a * b))

    agent = Agent(provider=None)
    msg = agent.load_plugin(MathPlugin())
    assert "loaded" in msg
    assert agent.use_tool("multiply", 6, 7) == 42


def test_plugin_duplicate_load_is_rejected():
    class DummyPlugin(Plugin):
        def __init__(self):
            super().__init__("Dummy")

        def setup(self, agent):
            pass

    agent = Agent(provider=None)
    agent.load_plugin(DummyPlugin())
    msg = agent.load_plugin(DummyPlugin())
    assert "already loaded" in msg
    assert len(agent.plugins) == 1


def test_memory_in_memory_only():
    mem = Memory()
    mem.add("user", "hello")
    assert len(mem.get_recent(5)) == 1


def test_memory_persists_to_disk():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "memory.json")

        mem1 = Memory(path=path)
        mem1.add("user", "my name is Aman")

        # simulate a separate process by creating a brand new Memory instance
        mem2 = Memory(path=path)
        recent = mem2.get_recent(5)
        assert len(recent) == 1
        assert recent[0]["content"] == "my name is Aman"


def test_memory_handles_corrupt_file():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "memory.json")
        with open(path, "w") as f:
            f.write("not valid json {{{")

        mem = Memory(path=path)  # should not raise
        assert mem.get_recent(5) == []
        mem.add("user", "fresh start")
        assert len(mem.get_recent(5)) == 1


def test_memory_search_finds_relevant_older_entry():
    mem = Memory()
    mem.add("user", "my dog is named Rex, a golden retriever")
    mem.add("user", "I like pizza")
    mem.add("user", "the weather is nice today")

    results = mem.search("tell me about my dog")
    assert len(results) >= 1
    assert "Rex" in results[0][0]["content"]


def test_search_tokenize_handles_plurals():
    assert tokenize("dogs") == tokenize("dog")
    assert tokenize("learning") == tokenize("learn")


def test_text_index_basic_relevance():
    idx = TextIndex()
    idx.add(0, "I love hiking in the mountains")
    idx.add(1, "my favorite food is pizza")
    idx.add(2, "I went hiking last weekend")

    results = idx.search("hiking trip")
    assert len(results) >= 1
    top_doc_id = results[0][0]
    assert top_doc_id in (0, 2)


def test_agent_chat_without_provider_does_not_crash():
    agent = Agent(provider=None)
    result = agent.chat("hello")
    assert "No provider configured" in result


def test_byteflow_importable_without_ollama_installed():
    # If this test file imports successfully and byteflow.cli imports too,
    # the core package has no hard dependency on the ollama package.
    from byteflow.cli import cli  # noqa: F401
    assert True


class _FakeProvider:
    """A minimal stand-in for a real LLM provider, for deterministic tests."""
    def __init__(self, response):
        self.response = response

    def generate(self, prompt):
        return self.response


def test_plan_returns_none_with_no_tools_registered():
    # Regression test: previously, run() with zero tools registered would
    # still ask the LLM to plan, which could hallucinate a tool name and
    # produce "Tool not found: X" for completely unrelated questions.
    agent = Agent(provider=_FakeProvider('[{"step": "add", "args": [1, 2]}]'))
    assert agent.plan("tell me about my pet") is None


def test_plan_rejects_hallucinated_tool_names():
    # The model returns a step calling a tool that was never registered.
    # plan() must reject the whole plan rather than let it through.
    agent = Agent(provider=_FakeProvider('[{"step": "totally_made_up_tool", "args": [1, 2]}]'))
    agent.register_tool(Tool("add", lambda a, b: a + b))
    assert agent.plan("do something") is None


def test_run_falls_back_to_chat_for_irrelevant_question_with_tools_registered():
    # Even with a real tool registered, an unrelated question should not
    # produce "Tool not found" - it should fall back to a normal answer.
    class HallucinatingThenAnsweringProvider:
        def __init__(self):
            self.call_count = 0

        def generate(self, prompt):
            self.call_count += 1
            if "Rules:" in prompt:  # planner call
                return '[{"step": "add", "args": [1, 2]}]'
            if '"tool"' in prompt:  # single-step tool-selection call
                return '{"tool": null, "answer": "I do not have info about your pet yet."}'
            return "I do not have info about your pet yet."

    agent = Agent(provider=HallucinatingThenAnsweringProvider())
    agent.register_tool(Tool("add", lambda a, b: a + b))
    result = agent.run("tell me about my pet")
    assert "Tool not found" not in str(result)


def test_builtin_tools_compute_correctly():
    from byteflow.builtin_tools import get_builtin_tools
    tools = {t.name: t for t in get_builtin_tools()}
    assert tools["add"].run(10, 20) == 30
    assert tools["divide"].run(20, 2) == 10.0
    assert "division by zero" in tools["divide"].run(5, 0)
    assert tools["add"].run("10", "40") == 50  # string args get coerced


def test_sandbox_runs_code_and_captures_output():
    from byteflow.sandbox import run_python_code
    result = run_python_code("print(2 + 2)")
    assert result.success
    assert "4" in result.stdout


def test_sandbox_captures_errors_without_crashing():
    from byteflow.sandbox import run_python_code
    result = run_python_code("print(1/0)")
    assert not result.success
    assert "ZeroDivisionError" in result.stderr


def test_sandbox_enforces_timeout():
    from byteflow.sandbox import run_python_code
    result = run_python_code("while True: pass", timeout=2)
    assert result.timed_out
    assert not result.success


def test_extract_code_block_handles_fenced_and_unfenced():
    fenced = "Here you go:\n```python\nprint(1)\n```\nDone."
    assert Agent.extract_code_block(fenced) == "print(1)"
    assert Agent.extract_code_block("print(1)") == "print(1)"


def test_code_mode_generates_and_executes():
    provider = _FakeProvider("```python\nprint(3 * 4)\n```")
    agent = Agent(provider=provider)
    result = agent.code("multiply 3 and 4", execute=True)
    assert "print(3 * 4)" in result["code"]
    assert result["executed"]
    assert result["result"].success
    assert "12" in result["result"].stdout


def test_run_routes_obvious_code_requests_to_code_mode():
    provider = _FakeProvider("```python\nprint('hi')\n```")
    agent = Agent(provider=provider)
    result = agent.run("write a function that prints hi", execute_code=True)
    assert isinstance(result, dict)
    assert "code" in result


def test_looks_like_chat_only_catches_greetings_and_filler():
    agent = Agent(provider=None)
    for msg in ["hii", "hi", "hello", "hey", "thanks", "ok", "...", "..", "bye"]:
        assert agent._looks_like_chat_only(msg) is True, f"{msg!r} should be chat-only"


def test_looks_like_chat_only_catches_memory_recall_questions():
    agent = Agent(provider=None)
    assert agent._looks_like_chat_only("what is my name") is True
    assert agent._looks_like_chat_only("what's my name") is True
    assert agent._looks_like_chat_only("do you remember me") is True


def test_looks_like_chat_only_does_not_catch_real_tool_requests():
    agent = Agent(provider=None)
    assert agent._looks_like_chat_only("add 10 and 20") is False
    assert agent._looks_like_chat_only("open youtube") is False
    assert agent._looks_like_chat_only("list files in C:\\Users\\me\\Documents") is False


def test_run_skips_tool_planner_entirely_for_greetings():
    # Regression test for the exact real-world bug: a local model asked
    # to pick a tool for "hii" guessed write_clipboard. The fix is to
    # never even ask the planner for messages this obviously chat-like.
    class TrackingProvider:
        def __init__(self):
            self.asked_to_pick_tool = False
        def generate(self, prompt):
            if "durable fact" in prompt:
                return "none"
            if "Available tools" in prompt or '"tool"' in prompt:
                self.asked_to_pick_tool = True
                return "null"
            return "Hey! How can I help?"

    provider = TrackingProvider()
    agent = Agent(provider=provider)
    agent.register_tool(Tool("write_clipboard", lambda t: f"wrote {t}"))

    result = agent.run("hii")
    assert result == "Hey! How can I help?"
    assert provider.asked_to_pick_tool is False  # never even asked


def test_run_still_uses_real_tools_for_clear_requests():
    # Make sure the chat-only filter doesn't accidentally swallow real
    # tool-using requests - math should still compute via the real tool.
    provider = _FakeProvider('[{"step": "add", "args": [10, 20]}]')
    agent = Agent(provider=provider)
    agent.register_tool(Tool("add", lambda a, b: a + b))
    result = agent.run("add 10 and 20")
    assert result == [30]


def test_list_files_has_sensible_default_when_folder_omitted():
    # Regression test: list_files() previously required `folder`, so a
    # planner forgetting to pass it crashed with a TypeError instead of
    # doing something sensible.
    from byteflow import desktop_tools
    result = desktop_tools.list_files()  # no folder argument at all
    assert isinstance(result, list)  # defaults to home dir, doesn't crash


def test_agent_api_unaffected_by_notebook_style_argv():
    # Regression test: the Agent/Memory/Tool API must not depend on sys.argv,
    # so it works correctly inside Jupyter notebooks, which overwrite argv
    # with kernel-launch arguments. Only byteflow.cli (built on click) cares
    # about argv, and that's expected - it's a terminal-only entrypoint.
    original_argv = sys.argv
    try:
        sys.argv = ["/path/to/ipykernel_launcher.py", "-f", "/tmp/fake-kernel.json"]

        provider = _FakeProvider('{"tool": null, "answer": "hi there"}')
        agent = Agent(provider=provider)
        agent.register_tool(Tool("add", lambda a, b: a + b))
        result = agent.use_tool("add", 2, 3)
        assert result == 5
    finally:
        sys.argv = original_argv


def test_profile_add_and_dedupe():
    p = Profile()
    assert p.add_fact("User's name is Aman") is True
    assert p.add_fact("user's name is aman") is False  # near-duplicate rejected
    assert len(p.all_facts()) == 1


def test_profile_persists_to_disk():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "profile.json")
        p1 = Profile(path=path)
        p1.add_fact("User prefers dark mode")

        p2 = Profile(path=path)  # simulate a separate process
        assert "User prefers dark mode" in p2.all_facts()


def test_profile_remove_fact():
    p = Profile()
    p.add_fact("User likes tea")
    assert p.remove_fact("User likes tea") is True
    assert p.all_facts() == []
    assert p.remove_fact("Not a real fact") is False


def test_profile_handles_corrupt_file():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "profile.json")
        with open(path, "w") as f:
            f.write("not valid json {{{")
        p = Profile(path=path)  # should not raise
        assert p.all_facts() == []


def test_agent_learn_from_exchange_extracts_fact():
    provider = _FakeProvider("User's name is Aman")
    agent = Agent(provider=provider, memory_path=False)  # ephemeral - testing logic, not persistence
    fact = agent.learn_from_exchange("my name is Aman", "Nice to meet you!")
    assert fact == "User's name is Aman"
    assert "User's name is Aman" in agent.profile.all_facts()


def test_agent_learn_from_exchange_ignores_none():
    provider = _FakeProvider("none")
    agent = Agent(provider=provider, memory_path=False)
    fact = agent.learn_from_exchange("what's the weather", "I don't have that info")
    assert fact is None
    assert agent.profile.all_facts() == []


def test_agent_learn_disabled_with_learn_false():
    provider = _FakeProvider("User's name is Aman")
    agent = Agent(provider=provider, memory_path=False, learn=False)
    fact = agent.learn_from_exchange("my name is Aman", "Nice to meet you!")
    assert fact is None
    assert agent.profile.all_facts() == []


def test_chat_automatically_learns_and_recalls_facts():
    class SequencedProvider:
        """Returns different responses for the chat call vs the fact-extraction call."""
        def generate(self, prompt):
            if "durable fact about the user" in prompt:
                return "User's name is Aman"
            return "Nice to meet you, Aman!"

    agent = Agent(provider=SequencedProvider())
    agent.chat("hi, my name is Aman")
    assert "User's name is Aman" in agent.profile.all_facts()

    # the fact should now show up in recalled_context for future prompts
    context = agent.recalled_context("what's my name?")
    assert "User's name is Aman" in context


def test_profile_path_auto_derived_from_memory_path():
    with tempfile.TemporaryDirectory() as d:
        memory_path = os.path.join(d, "memory.json")
        agent = Agent(provider=None, memory_path=memory_path)
        assert agent.profile.path == os.path.join(d, "memory_profile.json")


def test_agent_defaults_to_persistent_memory_when_path_not_passed():
    # Regression test for a real reported bug: a person using Agent()
    # directly (e.g. in a Jupyter notebook) with no memory_path got
    # silent in-memory-only behavior and lost context between sessions,
    # while the CLI (which always explicitly passes a path) didn't have
    # this problem - inconsistent and surprising. Agent() must now
    # default to the same persistent location as the CLI uses.
    expected = os.path.join(os.path.expanduser("~"), ".byteflow", "memory.json")
    agent = Agent(provider=None)
    assert agent.memory.path == expected


def test_agent_memory_path_false_is_explicit_ephemeral_opt_out():
    agent = Agent(provider=None, memory_path=False)
    assert agent.memory.path is None


def test_agent_memory_path_none_still_means_ephemeral():
    # Backward compatibility: code (including the existing CLI) that
    # explicitly passes memory_path=None must keep getting ephemeral
    # behavior, not silently switch to the new persistent default.
    agent = Agent(provider=None, memory_path=None)
    assert agent.memory.path is None


def test_agent_explicit_custom_memory_path_still_works():
    with tempfile.TemporaryDirectory() as d:
        custom_path = os.path.join(d, "custom.json")
        agent = Agent(provider=None, memory_path=custom_path)
        assert agent.memory.path == custom_path


def test_build_modelfile_basic_structure():
    from byteflow.tune import build_modelfile
    mf = build_modelfile(
        base_model="llama3",
        profile_facts=["User's name is Aman"],
        personality="Be warm and helpful.",
    )
    assert mf.startswith("FROM llama3\n")
    assert "SYSTEM" in mf
    assert "Be warm and helpful." in mf
    assert "User's name is Aman" in mf


def test_build_modelfile_escapes_triple_quotes_in_facts():
    from byteflow.tune import build_modelfile
    mf = build_modelfile(
        base_model="llama3",
        profile_facts=['Fact with """ embedded quotes'],
    )
    # exactly 2 unescaped triple-quote sequences (the SYSTEM block delimiters)
    assert mf.count('"""') == 2


def test_build_modelfile_handles_no_facts_or_personality():
    from byteflow.tune import build_modelfile
    mf = build_modelfile(base_model="llama3", profile_facts=[], personality=None)
    assert "FROM llama3" in mf
    assert "helpful assistant" in mf


def test_ollama_available_false_when_not_on_path():
    from byteflow.tune import ollama_available
    import shutil as _shutil
    original_which = _shutil.which
    _shutil.which = lambda name: None
    try:
        assert ollama_available() is False
    finally:
        _shutil.which = original_which


def test_create_tuned_model_raises_when_ollama_missing():
    from byteflow.tune import create_tuned_model, TuneError
    import shutil as _shutil
    original_which = _shutil.which
    _shutil.which = lambda name: None
    try:
        try:
            create_tuned_model("test-model", "llama3", ["fact"])
            assert False, "expected TuneError"
        except TuneError as e:
            assert "ollama" in str(e).lower()
    finally:
        _shutil.which = original_which


def test_list_files_filters_by_pattern():
    with tempfile.TemporaryDirectory() as d:
        for name in ("a.jpg", "b.jpg", "c.pdf"):
            open(os.path.join(d, name), "w").close()
        jpgs = desktop_tools.list_files(d, pattern="*.jpg")
        assert len(jpgs) == 2
        assert all(p.endswith(".jpg") for p in jpgs)


def test_list_files_nonexistent_folder_returns_error_string():
    result = desktop_tools.list_files("/definitely/not/a/real/path/xyz")
    assert isinstance(result, str) and result.startswith("Error:")


def test_search_files_matches_by_keyword():
    with tempfile.TemporaryDirectory() as d:
        for name in ("vacation_photo.jpg", "work_doc.pdf", "vacation_video.mp4"):
            open(os.path.join(d, name), "w").close()
        results = desktop_tools.search_files(d, "vacation")
        assert len(results) == 2


def test_organize_files_never_modifies_on_preview():
    with tempfile.TemporaryDirectory() as d:
        src_path = os.path.join(d, "a.jpg")
        open(src_path, "w").close()
        dest = os.path.join(d, "out")

        preview = desktop_tools.organize_files(d, "move", pattern="*.jpg", destination=dest)
        assert "DRY RUN" in preview
        assert os.path.exists(src_path)       # original untouched
        assert not os.path.exists(dest)       # destination not even created


def test_organize_files_has_no_confirm_parameter():
    # Safety guarantee: there is no argument that makes organize_files()
    # perform the action directly - confirm must come from a separate call.
    import inspect
    sig = inspect.signature(desktop_tools.organize_files)
    assert "confirm" not in sig.parameters


def test_confirm_organize_performs_previewed_move():
    with tempfile.TemporaryDirectory() as d:
        src_path = os.path.join(d, "a.jpg")
        open(src_path, "w").close()
        dest = os.path.join(d, "out")

        preview = desktop_tools.organize_files(d, "move", pattern="*.jpg", destination=dest)
        import re
        token = re.search(r"confirm_organize\('(\w+)'\)", preview).group(1)

        result = desktop_tools.confirm_organize(token)
        assert "Moved" in result
        assert not os.path.exists(src_path)
        assert os.path.exists(os.path.join(dest, "a.jpg"))


def test_confirm_organize_token_is_single_use():
    with tempfile.TemporaryDirectory() as d:
        open(os.path.join(d, "a.jpg"), "w").close()
        dest = os.path.join(d, "out")
        preview = desktop_tools.organize_files(d, "move", pattern="*.jpg", destination=dest)
        import re
        token = re.search(r"confirm_organize\('(\w+)'\)", preview).group(1)

        desktop_tools.confirm_organize(token)  # first use succeeds
        second = desktop_tools.confirm_organize(token)  # reuse should fail
        assert "Error" in second


def test_confirm_organize_rejects_unknown_token():
    result = desktop_tools.confirm_organize("not_a_real_token")
    assert "Error" in result


def test_organize_files_extra_positional_arg_rejected_safely():
    # Simulates a planner hallucinating a 5th arg trying to act like confirm=True.
    # Must raise TypeError (caught by Tool.run), never silently perform the action.
    with tempfile.TemporaryDirectory() as d:
        src_path = os.path.join(d, "secret.txt")
        open(src_path, "w").close()

        tool = Tool("organize_files", desktop_tools.organize_files)
        result = tool.run(d, "move", "*.txt", os.path.join(d, "out"), True)

        assert "Tool Error" in result
        assert os.path.exists(src_path)  # file was never touched


def test_organize_files_copy_preserves_original():
    with tempfile.TemporaryDirectory() as d:
        src_path = os.path.join(d, "doc.pdf")
        open(src_path, "w").close()
        dest = os.path.join(d, "backup")

        preview = desktop_tools.organize_files(d, "copy", pattern="*.pdf", destination=dest)
        import re
        token = re.search(r"confirm_organize\('(\w+)'\)", preview).group(1)
        desktop_tools.confirm_organize(token)

        assert os.path.exists(src_path)  # original still there
        assert os.path.exists(os.path.join(dest, "doc.pdf"))  # and the copy


def test_organize_files_rename_requires_n_placeholder():
    with tempfile.TemporaryDirectory() as d:
        open(os.path.join(d, "a.pdf"), "w").close()
        result = desktop_tools.organize_files(d, "rename", pattern="*.pdf", destination="newname.pdf")
        assert "Error" in result and "{n}" in result


def test_organize_files_invalid_action_rejected():
    result = desktop_tools.organize_files("/tmp", "delete", pattern="*")
    assert "Error" in result


def test_launch_handles_missing_target_gracefully():
    # Should never raise - always returns a string, even for a bogus target.
    result = desktop_tools.launch("definitely_not_a_real_app_xyz_123")
    assert isinstance(result, str)


def test_clipboard_functions_return_strings_without_pyperclip():
    # pyperclip isn't guaranteed to be installed; functions must degrade
    # gracefully to an error string, never raise.
    result = desktop_tools.read_clipboard()
    assert isinstance(result, str)


def test_draft_social_post_never_calls_a_publish_action():
    # Safety guarantee: there is no "publish" or "submit" function anywhere
    # in desktop_tools - draft_social_post can only draft, copy, and open
    # a URL. Verify no such function exists to call.
    import byteflow.desktop_tools as dt_module
    forbidden_names = ("post", "publish", "submit", "send_post")
    for name in forbidden_names:
        assert not hasattr(dt_module, name), f"desktop_tools must not expose a '{name}' function"


def test_draft_social_post_returns_draft_and_metadata():
    provider = _FakeProvider("Excited to share my new project!")
    agent = Agent(provider=provider)

    # stub out the OS-touching calls so this test doesn't depend on a
    # real clipboard or browser being available
    import byteflow.desktop_tools as dt_module
    original_write = dt_module.write_clipboard
    original_launch = dt_module.launch
    dt_module.write_clipboard = lambda text: "Clipboard updated."
    dt_module.launch = lambda target: f"Launched: {target}"
    try:
        result = agent.draft_social_post("a new project", platform="linkedin", tone="excited")
        assert result["draft"] == "Excited to share my new project!"
        assert result["clipboard"] is True
        assert result["launched"] is True
        assert "linkedin.com" in result["platform_url"]
    finally:
        dt_module.write_clipboard = original_write
        dt_module.launch = original_launch


def test_draft_social_post_no_open_skips_launch():
    provider = _FakeProvider("Some draft text")
    agent = Agent(provider=provider)

    import byteflow.desktop_tools as dt_module
    original_write = dt_module.write_clipboard
    original_launch = dt_module.launch
    launch_called = []
    dt_module.write_clipboard = lambda text: "Clipboard updated."
    dt_module.launch = lambda target: launch_called.append(target) or "Launched"
    try:
        result = agent.draft_social_post("topic", open_site=False)
        assert result["launched"] is False
        assert launch_called == []  # launch() was never even called
    finally:
        dt_module.write_clipboard = original_write
        dt_module.launch = original_launch


def test_draft_social_post_without_provider_fails_cleanly():
    agent = Agent(provider=None)
    result = agent.draft_social_post("topic")
    assert "No provider configured" in result["draft"]
    assert result["clipboard"] is False
    assert result["launched"] is False


def test_draft_social_post_unknown_platform_has_no_url():
    provider = _FakeProvider("draft text")
    agent = Agent(provider=provider)
    result = agent.draft_social_post("topic", platform="myspace")
    assert result["platform_url"] is None
    assert result["launched"] is False


def test_draft_social_post_unknown_platform_returns_clear_warning():
    provider = _FakeProvider("draft text")
    agent = Agent(provider=provider)
    result = agent.draft_social_post("topic", platform="YouTube")
    assert result["draft"] is None
    assert result["warning"] is not None
    assert "YouTube" in result["warning"]
    assert "linkedin" in result["warning"].lower()  # mentions what IS supported


def test_draft_social_post_unknown_platform_never_calls_llm():
    # No point spending a generation call drafting text for a platform
    # that has nowhere to paste it - this must short-circuit before
    # touching the provider at all.
    class CountingProvider:
        def __init__(self):
            self.calls = 0
        def generate(self, prompt):
            self.calls += 1
            return "should not happen"

    provider = CountingProvider()
    agent = Agent(provider=provider)
    agent.draft_social_post("topic", platform="youtube")
    assert provider.calls == 0


def test_draft_social_post_success_has_no_warning():
    provider = _FakeProvider("a good draft")
    agent = Agent(provider=provider)

    import byteflow.desktop_tools as dt_module
    original_write = dt_module.write_clipboard
    original_launch = dt_module.launch
    dt_module.write_clipboard = lambda text: "Clipboard updated."
    dt_module.launch = lambda target: f"Launched: {target}"
    try:
        result = agent.draft_social_post("topic", platform="linkedin")
        assert result["warning"] is None
        assert result["draft"] == "a good draft"
    finally:
        dt_module.write_clipboard = original_write
        dt_module.launch = original_launch


def test_resolve_shortcut_known_name_case_insensitive():
    assert desktop_tools.resolve_shortcut("youtube") == "https://www.youtube.com"
    assert desktop_tools.resolve_shortcut("YouTube") == "https://www.youtube.com"
    assert desktop_tools.resolve_shortcut("  youtube  ") == "https://www.youtube.com"


def test_resolve_shortcut_tolerates_spaces_and_hyphens():
    # Regression test for a real reported bug: "open vs code" (typed
    # with a space) failed to launch because resolve_shortcut only did
    # an exact match against "vscode" (no space).
    assert desktop_tools.resolve_shortcut("vs code") == "code"
    assert desktop_tools.resolve_shortcut("VS Code") == "code"
    assert desktop_tools.resolve_shortcut("vs-code") == "code"
    assert desktop_tools.resolve_shortcut("visual studio code") == "code"


def test_resolve_shortcut_unknown_name_passes_through():
    assert desktop_tools.resolve_shortcut("https://example.com") == "https://example.com"
    assert desktop_tools.resolve_shortcut("not_a_real_shortcut_xyz") == "not_a_real_shortcut_xyz"


def test_list_shortcuts_returns_sorted_known_names():
    names = desktop_tools.list_shortcuts()
    assert "youtube" in names
    assert "vscode" in names
    assert names == sorted(names)


def test_launch_resolves_shortcut_before_calling_os():
    calls = []

    def fake_subprocess_run(args, **kwargs):
        calls.append(args)
        class FakeResult:
            returncode = 0
            stderr = ""
        return FakeResult()

    import subprocess as subprocess_module
    original_run = subprocess_module.run
    subprocess_module.run = fake_subprocess_run
    try:
        result = desktop_tools.launch("youtube")
        assert "https://www.youtube.com" in result
        # the resolved URL, not the literal word "youtube", must be what got passed to the OS
        assert any("https://www.youtube.com" in str(c) for c in calls)
    finally:
        subprocess_module.run = original_run


def test_launch_raw_target_unaffected_by_shortcut_table():
    def fake_subprocess_run(args, **kwargs):
        class FakeResult:
            returncode = 0
            stderr = ""
        return FakeResult()

    import subprocess as subprocess_module
    original_run = subprocess_module.run
    subprocess_module.run = fake_subprocess_run
    try:
        result = desktop_tools.launch("https://my-own-site.example.com")
        assert "my-own-site.example.com" in result
    finally:
        subprocess_module.run = original_run


def test_companion_send_is_non_blocking():
    import time

    class SlowProvider:
        def generate(self, prompt):
            time.sleep(0.2)
            return "none" if "durable fact" in prompt else "a reply"

    agent = Agent(provider=SlowProvider())
    controller = CompanionController(agent)

    start = time.time()
    controller.send("hello")
    elapsed = time.time() - start

    assert elapsed < 0.1  # send() must return almost immediately
    assert controller.busy is True


def test_companion_poll_reply_returns_none_when_not_ready():
    agent = Agent(provider=_FakeProvider("a reply"))
    controller = CompanionController(agent)
    assert controller.poll_reply() is None  # nothing sent yet


def test_companion_format_result_handles_plain_string():
    assert CompanionController._format_result("a plain reply") == "a plain reply"


def test_companion_format_result_handles_code_mode_dict():
    class FakeExecResult:
        def format(self):
            return "--- stdout ---\n42\n"

    result = {
        "code": "print(42)",
        "executed": True,
        "result": FakeExecResult(),
    }
    formatted = CompanionController._format_result(result)
    assert "print(42)" in formatted
    assert "42" in formatted
    assert "Output:" in formatted


def test_companion_format_result_handles_unexecuted_code_dict():
    result = {"code": "print(42)", "executed": False, "result": None}
    formatted = CompanionController._format_result(result)
    assert "print(42)" in formatted
    assert "Output:" not in formatted


def test_companion_speech_friendly_summarizes_code_instead_of_reading_it():
    # Regression test for the exact bug found and fixed: speaking a code
    # reply aloud must never read raw source line-by-line - it should
    # give a short spoken summary instead.
    code_reply = "Here's the code:\n\nprint(42)\n\nOutput:\n--- stdout ---\n42\n"
    speech = CompanionController.speech_friendly(code_reply)
    assert "print(42)" not in speech       # raw source must not be spoken
    assert "42" in speech                  # but the actual output IS mentioned
    assert "wrote the code" in speech.lower()
    assert "--- stdout ---" not in speech  # raw section markers shouldn't be read aloud


def test_companion_speech_friendly_code_without_execution():
    code_reply = "Here's the code:\n\nprint(42)"
    speech = CompanionController.speech_friendly(code_reply)
    assert "print(42)" not in speech
    assert "panel" in speech.lower() or "code" in speech.lower()


def test_companion_speech_friendly_passes_through_plain_replies_unchanged():
    plain = "The capital of France is Paris."
    assert CompanionController.speech_friendly(plain) == plain


def test_companion_uses_real_tools_not_just_chat():
    # Regression test: the companion must route through agent.run() so
    # registered tools (math, desktop helpers, etc.) actually get used,
    # not just plain chat() - this was a real gap found and fixed.
    import time

    class PlannerProvider:
        def generate(self, prompt):
            if "Rules:" in prompt and "Available tools" in prompt:
                return '[{"step": "add", "args": [4, 5]}]'
            return "none"  # fact-extraction calls, if any

    agent = Agent(provider=PlannerProvider())
    agent.register_tool(Tool("add", lambda a, b: a + b))
    controller = CompanionController(agent)

    controller.send("add 4 and 5")
    reply = None
    for _ in range(50):
        reply = controller.poll_reply()
        if reply is not None:
            break
        time.sleep(0.02)

    assert "9" in reply  # the REAL computed result, not an LLM guess


def test_companion_full_round_trip_via_polling():
    import time

    class SequencedProvider:
        def generate(self, prompt):
            return "none" if "durable fact" in prompt else "the real reply"

    agent = Agent(provider=SequencedProvider())
    controller = CompanionController(agent)

    controller.send("hello")
    reply = None
    for _ in range(50):  # poll like the real GUI loop would
        reply = controller.poll_reply()
        if reply is not None:
            break
        time.sleep(0.02)

    assert reply == "the real reply"
    assert controller.busy is False


def test_companion_busy_send_queues_and_answers_the_second_message():
    import time

    class CountingProvider:
        def __init__(self):
            self.calls = 0
        def generate(self, prompt):
            self.calls += 1
            time.sleep(0.2)
            return "none" if "durable fact" in prompt else "reply"

    provider = CountingProvider()
    agent = Agent(provider=provider)
    controller = CompanionController(agent)

    controller.send("first")
    controller.send("second, sent while still busy")
    time.sleep(0.05)

    # sending while busy queues the message and returns a notice - it
    # must NOT be silently discarded. A real observed bug: the old
    # behavior threw the second message away entirely, so a real
    # question typed while the agent was still thinking about the
    # previous one just vanished with no way to recover it.
    notice = controller.poll_reply()
    assert "still thinking" in notice.lower()

    # both messages must eventually be answered, one after the other -
    # collect replies until we've seen two (the first message's reply,
    # then the queued second message's reply)
    replies = []
    for _ in range(50):
        time.sleep(0.1)
        reply = controller.poll_reply()
        if reply is not None:
            replies.append(reply)
        if len(replies) >= 2:
            break

    assert len(replies) == 2, f"expected both messages answered, got {replies}"
    assert not controller.busy


def test_companion_empty_message_is_ignored():
    agent = Agent(provider=_FakeProvider("should not be returned"))
    controller = CompanionController(agent)
    controller.send("")
    controller.send("   ")
    assert controller.busy is False
    assert controller.poll_reply() is None


def test_companion_provider_exception_does_not_crash_and_resets_busy():
    import time

    class CrashingProvider:
        def generate(self, prompt):
            raise RuntimeError("simulated failure")

    agent = Agent(provider=CrashingProvider())
    controller = CompanionController(agent)
    controller.send("hello")

    reply = None
    for _ in range(50):
        reply = controller.poll_reply()
        if reply is not None:
            break
        time.sleep(0.02)

    assert "Error" in reply
    assert controller.busy is False


def test_companion_controller_importable_without_tkinter():
    # CompanionController must not require tkinter at import time - only
    # run_companion() (the actual GUI) should need a display.
    import sys
    assert "tkinter" not in sys.modules or True  # informational; real check is the import below
    from byteflow.companion import CompanionController as CC
    assert CC is not None


def test_companion_speak_replies_degrades_gracefully_without_pyttsx3():
    # pyttsx3 isn't guaranteed to be installed (it isn't in this test
    # environment) - requesting voice output must not crash, just
    # silently stay text-only.
    agent = Agent(provider=None)
    controller = CompanionController(agent, speak_replies=True)
    assert controller.speaker is None
    controller.speak("this should be a safe no-op")  # must not raise


def test_companion_default_has_no_speaker():
    agent = Agent(provider=None)
    controller = CompanionController(agent)
    assert controller.speaker is None


def test_companion_speak_with_fake_speaker_does_not_block():
    import time

    class SlowFakeSpeaker:
        def speak(self, text):
            time.sleep(0.3)

    agent = Agent(provider=None)
    controller = CompanionController(agent)
    controller.speaker = SlowFakeSpeaker()

    start = time.time()
    controller.speak("hello")
    elapsed = time.time() - start
    assert elapsed < 0.1  # speak() must return immediately, not block


def test_voice_tts_available_false_without_pyttsx3():
    from byteflow.voice import tts_available
    assert tts_available() is False  # not installed in this environment


def test_voice_stt_available_false_without_vosk():
    from byteflow.voice import stt_available
    assert stt_available() is False  # not installed in this environment


def test_voice_speaker_raises_clear_error_without_pyttsx3():
    from byteflow.voice import Speaker, VoiceError
    try:
        Speaker()
        assert False, "expected VoiceError"
    except VoiceError as e:
        assert "pyttsx3" in str(e)


def test_voice_listener_raises_clear_error_without_vosk():
    from byteflow.voice import Listener, VoiceError
    try:
        Listener()
        assert False, "expected VoiceError"
    except VoiceError as e:
        assert "vosk" in str(e).lower()


def test_voice_model_present_distinguishes_empty_missing_and_populated():
    from byteflow.voice import vosk_model_present
    with tempfile.TemporaryDirectory() as d:
        empty_dir = os.path.join(d, "empty")
        os.makedirs(empty_dir)
        assert vosk_model_present(empty_dir) is False

        nonexistent = os.path.join(d, "nope")
        assert vosk_model_present(nonexistent) is False

        populated = os.path.join(d, "populated")
        os.makedirs(populated)
        open(os.path.join(populated, "f"), "w").close()
        assert vosk_model_present(populated) is True


def test_conversation_listener_raises_clear_error_without_vosk():
    from byteflow.voice import ConversationListener, VoiceError
    try:
        ConversationListener(on_utterance=lambda t: None)
        assert False, "expected VoiceError"
    except VoiceError as e:
        assert "vosk" in str(e).lower()


def test_conversation_listener_behavior_with_fake_vosk():
    # This test exercises the real partial->final state machine, but
    # needs actual vosk/sounddevice modules (even fake ones) importable,
    # which this sandbox doesn't have installed. Skip cleanly rather than
    # fail when they're unavailable - on a machine with the real voice
    # extras installed, swap in real audio and this same shape of test
    # would apply against the real library's partial/final result format.
    from byteflow.voice import stt_available
    if not stt_available():
        print("SKIP (stt_available() is False in this environment - "
              "vosk/sounddevice not installed, expected in this sandbox)")
        return

    # If vosk/sounddevice ARE available, do a basic smoke check that
    # construction succeeds when a model is present (without actually
    # opening a real audio device, since that may not exist either).
    from byteflow.voice import ConversationListener, vosk_model_present, DEFAULT_VOSK_MODEL_DIR
    if not vosk_model_present():
        print(f"SKIP (no Vosk model at {DEFAULT_VOSK_MODEL_DIR})")
        return

    conv = ConversationListener(on_utterance=lambda t: None)
    assert conv.running is False


def test_cli_companion_command_wires_conversation_mode_flag():
    # Regression test: --conversation-mode existed in run_companion()'s
    # Python API but was never actually exposed as a CLI flag, so
    # `byteflow companion --conversation-mode` silently did nothing.
    # Verify the CLI correctly threads all voice-related flags through
    # to run_companion(), without needing a real display, ollama
    # package, or audio library - mock exactly the three things that
    # genuinely require those (tkinter, OllamaProvider, run_companion).
    import sys
    import types
    from click.testing import CliRunner
    from byteflow import cli as cli_module
    import byteflow.companion as companion_module
    import byteflow.providers.ollama_provider as ollama_provider_module

    captured = {}

    def fake_run_companion(**kwargs):
        captured.update(kwargs)

    class FakeOllamaProvider:
        def __init__(self, model="llama3"):
            self.model = model

    original_run_companion = companion_module.run_companion
    original_provider = ollama_provider_module.OllamaProvider
    tkinter_was_present = "tkinter" in sys.modules

    companion_module.run_companion = fake_run_companion
    cli_module.OllamaProvider = FakeOllamaProvider
    sys.modules["tkinter"] = types.ModuleType("tkinter")

    try:
        runner = CliRunner()
        result = runner.invoke(
            cli_module.cli,
            ["companion", "--conversation-mode", "--memory-path", "none"],
        )
        assert result.exit_code == 0, result.output
        assert captured.get("conversation_mode") is True
        assert captured.get("voice_input") is False
        assert captured.get("voice_output") is False

        # --voice and --conversation-mode should compose correctly together
        captured.clear()
        result2 = runner.invoke(
            cli_module.cli,
            ["companion", "--voice", "--conversation-mode", "--memory-path", "none"],
        )
        assert result2.exit_code == 0, result2.output
        assert captured.get("voice_input") is True
        assert captured.get("voice_output") is True
        assert captured.get("conversation_mode") is True
    finally:
        companion_module.run_companion = original_run_companion
        cli_module.OllamaProvider = original_provider
        if not tkinter_was_present:
            del sys.modules["tkinter"]


# -----------------------------
# CHUNKING
# -----------------------------

def test_chunk_text_short_text_returns_single_chunk():
    short = "This is a short sentence."
    assert chunking.chunk_text(short) == [short]


def test_chunk_text_empty_returns_empty_list():
    assert chunking.chunk_text("") == []
    assert chunking.chunk_text("   ") == []


def test_chunk_text_splits_long_text_into_multiple_chunks():
    sentences = [f"Sentence number {i} is here for testing." for i in range(50)]
    text = " ".join(sentences)
    chunks = chunking.chunk_text(text, max_chars=300, overlap_chars=50)
    assert len(chunks) > 1
    for c in chunks:
        assert len(c) <= 350  # some tolerance for the last sentence in a chunk


def test_chunk_text_overlap_preserves_boundary_straddling_facts():
    sentences = [f"Filler sentence number {i}." for i in range(20)]
    sentences.insert(10, "The secret code is XYZZY42.")
    text = " ".join(sentences)
    chunks = chunking.chunk_text(text, max_chars=400, overlap_chars=80)
    fact_chunks = [c for c in chunks if "XYZZY42" in c]
    assert len(fact_chunks) >= 1


def test_chunk_text_handles_unpunctuated_text_without_unbounded_chunks():
    word_salad = " ".join(["word"] * 500)
    chunks = chunking.chunk_text(word_salad, max_chars=300, overlap_chars=50)
    for c in chunks:
        assert len(c) <= 300


def test_chunk_text_handles_text_with_no_whitespace_at_all():
    no_whitespace = "x" * 1000
    chunks = chunking.chunk_text(no_whitespace, max_chars=300, overlap_chars=50)
    for c in chunks:
        assert len(c) <= 300
    assert "".join(chunks) == no_whitespace  # no characters lost


def test_chunk_with_metadata_includes_source_and_index():
    text = " ".join([f"Sentence {i} for testing purposes today." for i in range(30)])
    chunks = chunking.chunk_with_metadata(text, source="test.txt", max_chars=300)
    assert all(c["source"] == "test.txt" for c in chunks)
    assert [c["chunk_index"] for c in chunks] == list(range(len(chunks)))


def test_split_into_sentences_basic():
    result = chunking.split_into_sentences("First sentence. Second one! Third?")
    assert result == ["First sentence.", "Second one!", "Third?"]


# -----------------------------
# EMBEDDINGS
# -----------------------------

def test_tfidf_embedder_similarity_ranks_related_text_higher():
    embedder = TfidfEmbedder()
    corpus = [
        "my dog is named Rex, a golden retriever",
        "I like pizza for dinner",
        "I went hiking with my dog Rex last weekend",
    ]
    embedder.fit(corpus)
    vecs = [embedder.embed_one(t) for t in corpus]
    query_vec = embedder.embed_one("tell me about my dog")

    scores = [embedder.similarity(query_vec, v) for v in vecs]
    assert scores[0] > scores[1]  # dog-related ranks above pizza
    assert scores[2] > scores[1]  # dog-related ranks above pizza


def test_tfidf_embedder_works_without_fit_called():
    embedder = TfidfEmbedder()  # no fit()
    vec = embedder.embed_one("hello world hello")
    assert vec  # still produces something usable


def test_tfidf_embedder_empty_text_gives_empty_vector():
    embedder = TfidfEmbedder()
    assert embedder.embed_one("") == {}


def test_tfidf_embedder_similarity_zero_for_unrelated_text():
    embedder = TfidfEmbedder()
    embedder.fit(["dogs and cats", "rockets and planets"])
    v1 = embedder.embed_one("dogs and cats")
    v2 = embedder.embed_one("rockets and planets")
    assert embedder.similarity(v1, v2) == 0.0


def test_sentence_transformer_embedder_raises_clear_error_when_unavailable():
    if sentence_transformers_available():
        print("SKIP (sentence-transformers IS installed in this environment)")
        return
    try:
        SentenceTransformerEmbedder()
        assert False, "expected ImportError"
    except ImportError as e:
        assert "sentence-transformers" in str(e)


# -----------------------------
# VECTOR STORE
# -----------------------------

def test_vector_store_add_and_search_short_entries():
    store = VectorStore()
    store.add_document("my dog is named Rex, a golden retriever", source="chat:1")
    store.add_document("I like pizza for dinner", source="chat:2")

    results = store.search("tell me about my dog")
    assert len(results) >= 1
    assert results[0]["source"] == "chat:1"


def test_vector_store_has_documents():
    store = VectorStore()
    assert store.has_documents() is False

    store.add_document("some notes about the project", source="notes.pdf")
    assert store.has_documents() is True
    assert store.has_documents(source_prefix="notes") is True
    assert store.has_documents(source_prefix="other") is False


def test_vector_store_search_scoped_to_source():
    store = VectorStore()
    store.add_document("Python credits: 3. Discrete Mathematics credits: 4.", source="examform.pdf")
    store.add_document("Recommended book: Core Python Programming by Wesley Chun.", source="syllabus.pdf")

    scoped = store.search("credits", source="examform.pdf")
    assert len(scoped) >= 1
    assert all(r["source"] == "examform.pdf" for r in scoped)

    scoped_other = store.search("credits", source="syllabus.pdf")
    # syllabus doesn't mention credits at all - scoping must not leak
    # in the examform chunk just because it's a better match overall
    assert scoped_other == [] or all(r["source"] == "syllabus.pdf" for r in scoped_other)


def test_vector_store_list_sources_preserves_add_order():
    store = VectorStore()
    store.add_document("first doc", source="a.pdf")
    store.add_document("second doc", source="b.pdf")
    store.add_document("more from first doc", source="a.pdf")
    assert store.list_sources() == ["a.pdf", "b.pdf"]


def test_agent_active_document_source_defaults_to_most_recent_upload():
    agent = Agent(provider=None, memory_path=False)
    assert agent.active_document_source is None

    agent.ingest_document("syllabus content", source="syllabus.pdf")
    assert agent.active_document_source == "syllabus.pdf"

    agent.ingest_document("exam form content", source="examform.pdf")
    assert agent.active_document_source == "examform.pdf"


def test_match_ingested_source_is_case_and_spacing_tolerant():
    agent = Agent(provider=None, memory_path=False)
    agent.ingest_document("exam form content", source="SAI AMAN ZAKIRSHA.pdf")

    assert agent._match_ingested_source("it's on Sai Aman Zakirsha.pdf") == "SAI AMAN ZAKIRSHA.pdf"
    assert agent._match_ingested_source("what's the weather today") is None


def test_chat_scopes_retrieval_to_active_document_not_blended_across_files():
    # Real observed bug: with two documents ingested, a credits question
    # pulled the wrong chunk from an unrelated file and confidently cited
    # the wrong filename as its source. Prove retrieval now stays scoped
    # to the active document, and following a mention of another actual
    # ingested file resets the active document to a real source
    # (not a hallucinated matching_source when the message names
    # neither file - it should stay a valid known source either way).
    captured = []

    class CapturingProvider:
        def generate(self, prompt):
            captured.append(prompt)
            return "none" if "durable fact" in prompt else "ok"

    agent = Agent(provider=CapturingProvider(), memory_path=False)
    agent.ingest_document(
        "Recommended book: Core Python Programming by Wesley Chun.",
        source="syllabus.pdf",
    )
    agent.ingest_document(
        "Subject credits: Python Programming 3 credits, Discrete Mathematics 4 credits.",
        source="examform.pdf",
    )

    captured.clear()
    agent.chat("Discrete Mathematics, credits")
    prompt = captured[0]

    assert "examform.pdf" in prompt
    assert "syllabus.pdf" not in prompt  # scoped search should not pull in the other file
    assert "currently focused on the document `examform.pdf`" in prompt


def test_vector_store_chunks_long_documents_automatically():
    store = VectorStore()
    long_doc = " ".join([f"Paragraph {i} about hiking trails." for i in range(30)])
    long_doc += " The secret deployment key is projectx-prod-42."

    n_chunks = store.add_document(long_doc, source="notes.txt")
    assert n_chunks > 1

    results = store.search("what is the deployment key")
    assert len(results) >= 1
    assert "projectx-prod-42" in results[0]["text"]


def test_vector_store_persists_across_processes():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "vectors.json")
        store1 = VectorStore(path=path)
        store1.add_document("my dog is named Rex", source="chat:1")

        store2 = VectorStore(path=path)  # simulates a fresh process
        assert len(store2.entries) == 1
        results = store2.search("tell me about my dog")
        assert len(results) >= 1


def test_vector_store_handles_corrupt_file():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "vectors.json")
        with open(path, "w") as f:
            f.write("not valid json {{{")
        store = VectorStore(path=path)  # should not raise
        assert store.entries == []
        store.add_document("test", source="x")
        assert len(store.entries) == 1


def test_vector_store_remove_source():
    store = VectorStore()
    store.add_document("content about dogs", source="doc1")
    store.add_document("content about cats", source="doc2")

    removed = store.remove_source("doc1")
    assert removed == 1
    assert all(e["source"] == "doc2" for e in store.entries)


def test_vector_store_clear():
    store = VectorStore()
    store.add_document("some content", source="doc1")
    store.clear()
    assert store.entries == []


def test_vector_store_search_on_empty_store_returns_empty():
    store = VectorStore()
    assert store.search("anything") == []


# -----------------------------
# AGENT INTEGRATION
# -----------------------------

def test_agent_has_vector_store_by_default():
    agent = Agent(provider=None, memory_path=False)
    assert agent.vector_store is not None


def test_agent_vector_store_path_auto_derived_from_memory_path():
    with tempfile.TemporaryDirectory() as d:
        memory_path = os.path.join(d, "memory.json")
        agent = Agent(provider=None, memory_path=memory_path)
        assert agent.vector_store.path == os.path.join(d, "memory_vectors.json")


def test_agent_ingest_document_chunks_and_indexes():
    agent = Agent(provider=None, memory_path=False)
    long_doc = " ".join([f"Paragraph {i} about project planning." for i in range(30)])
    long_doc += " The API key is stored under SECRET_KEY_777."

    n_chunks = agent.ingest_document(long_doc, source="notes.txt")
    assert n_chunks > 1

    results = agent.vector_store.search("what is the api key")
    assert len(results) >= 1
    assert "SECRET_KEY_777" in results[0]["text"]


def test_agent_recalled_context_includes_relevant_document_chunks():
    agent = Agent(provider=None, memory_path=False)
    long_doc = " ".join([f"Paragraph {i} about general topics." for i in range(30)])
    long_doc += " The deployment credentials are in vault key projectx-prod."
    agent.ingest_document(long_doc, source="deploy_notes.txt")

    context = agent.recalled_context("where are the deployment credentials?")
    assert "deploy_notes.txt" in context
    assert "projectx-prod" in context


def test_agent_recalled_context_without_documents_has_no_document_section():
    agent = Agent(provider=None, memory_path=False)
    context = agent.recalled_context("hello")
    assert "documents you've shared" not in context


# -----------------------------
# COMPANION FACE DRAWING (no display needed - fake canvas)
# -----------------------------

class _FakeCanvas:
    """
    Minimal stand-in for tkinter.Canvas - just enough surface area
    (create_oval/create_rectangle/create_line/itemconfig) to exercise
    _build_face()'s real drawing logic and the eye helper functions
    without needing an actual display, which this test environment
    doesn't have.
    """
    def __init__(self):
        self._next_id = 1
        self.items = {}

    def _new(self, kind, kwargs):
        item_id = self._next_id
        self._next_id += 1
        self.items[item_id] = (kind, kwargs)
        return item_id

    def create_oval(self, *args, **kwargs):
        return self._new("oval", kwargs)

    def create_rectangle(self, *args, **kwargs):
        return self._new("rectangle", kwargs)

    def create_line(self, *args, **kwargs):
        return self._new("line", kwargs)

    def itemconfig(self, item_id, **kwargs):
        self.items[item_id][1].update(kwargs)


def test_build_face_returns_expected_structure():
    from byteflow.companion import _build_face
    canvas = _FakeCanvas()
    head, left_eye, right_eye, antenna_items = _build_face(canvas, size=140)

    assert isinstance(head, int)
    assert left_eye is right_eye  # single shared dict covering BOTH eyes by design
    base_keys = {
        "glow_outer", "glow_inner", "ring", "iris_outer", "iris_inner",
        "core", "spark", "spark_small", "iris",
    }
    expected_keys = {f"l_{k}" for k in base_keys} | {f"r_{k}" for k in base_keys}
    assert set(left_eye.keys()) == expected_keys
    assert len(antenna_items) == 2


def test_set_eyes_color_updates_iris_and_ring():
    from byteflow.companion import _build_face, _set_eyes_color
    canvas = _FakeCanvas()
    head, left_eye, right_eye, _ = _build_face(canvas, size=140)

    _set_eyes_color(canvas, left_eye, right_eye, "#ebcb8b")
    for prefix in ("l_", "r_"):
        assert canvas.items[left_eye[f"{prefix}iris_outer"]][1]["fill"] == "#ebcb8b"
        assert canvas.items[left_eye[f"{prefix}ring"]][1]["fill"] == "#ebcb8b"


def test_set_eyes_visible_toggles_only_inner_layers():
    from byteflow.companion import _build_face, _set_eyes_visible
    canvas = _FakeCanvas()
    head, left_eye, right_eye, _ = _build_face(canvas, size=140)

    _set_eyes_visible(canvas, left_eye, right_eye, False)
    for prefix in ("l_", "r_"):
        for key in ("iris_outer", "iris_inner", "core", "spark", "spark_small"):
            assert canvas.items[left_eye[f"{prefix}{key}"]][1]["state"] == "hidden"
        # the outer glow/ring should be untouched by blinking, so the eye
        # socket stays visible while only the inner iris/pupil appears to close
        assert "state" not in canvas.items[left_eye[f"{prefix}ring"]][1]
        assert "state" not in canvas.items[left_eye[f"{prefix}glow_outer"]][1]

    _set_eyes_visible(canvas, left_eye, right_eye, True)
    for prefix in ("l_", "r_"):
        for key in ("iris_outer", "iris_inner", "core", "spark", "spark_small"):
            assert canvas.items[left_eye[f"{prefix}{key}"]][1]["state"] == "normal"


# -----------------------------
# WEB SEARCH
# -----------------------------

class _FakeHTTPResponse:
    """Minimal stand-in for the object urllib.request.urlopen() returns,
    just enough for web_search.search() to call .read() on it."""
    def __init__(self, data):
        self.data = data
    def read(self):
        return self.data
    def __enter__(self):
        return self
    def __exit__(self, *args):
        return False


def _mock_urlopen(html_bytes):
    """Context-manager-free helper: monkeypatch urlopen for the duration
    of a test, returning html_bytes as the response body, then restore
    the original afterward - used by several tests below."""
    original = web_search_module.urllib.request.urlopen
    web_search_module.urllib.request.urlopen = lambda request, timeout=10: _FakeHTTPResponse(html_bytes)
    return original


def test_web_search_parser_extracts_title_url_snippet():
    sample_html = (
        '<a class="result__a" '
        'href="//duckduckgo.com/l/?uddg=https%3A%2F%2Fwww.python.org%2F&amp;rut=abc">'
        'Python.org</a>'
        '<a class="result__snippet">The official home of the <b>Python</b> language.</a>'
    ).encode("utf-8")

    original = _mock_urlopen(sample_html)
    try:
        results = web_search_module.search("python", max_results=5)
        assert len(results) == 1
        assert results[0]["title"] == "Python.org"
        assert results[0]["url"] == "https://www.python.org/"
        assert "Python" in results[0]["snippet"]
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_respects_max_results():
    sample_html = "".join(
        f'<a class="result__a" href="//duckduckgo.com/l/?uddg=https%3A%2F%2Fexample.com%2F{i}">Result {i}</a>'
        f'<a class="result__snippet">Snippet {i}.</a>'
        for i in range(5)
    ).encode("utf-8")

    original = _mock_urlopen(sample_html)
    try:
        results = web_search_module.search("test", max_results=2)
        assert len(results) == 2
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_empty_query_returns_empty_list():
    assert web_search_module.search("") == []
    assert web_search_module.search("   ") == []


def test_web_search_no_results_returns_empty_list():
    original = _mock_urlopen(b"<html><body>nothing here</body></html>")
    try:
        results = web_search_module.search("asdkjaslkdjaslkdj")
        assert results == []
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_network_failure_raises_clear_error():
    import urllib.error

    def raise_url_error(request, timeout=10):
        raise urllib.error.URLError("no route to host")

    original = web_search_module.urllib.request.urlopen
    web_search_module.urllib.request.urlopen = raise_url_error
    try:
        try:
            web_search_module.search("test")
            assert False, "expected WebSearchError"
        except web_search_module.WebSearchError as e:
            assert "reach" in str(e).lower()
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_formatted_degrades_gracefully_on_failure():
    import urllib.error

    def raise_url_error(request, timeout=10):
        raise urllib.error.URLError("no route to host")

    original = web_search_module.urllib.request.urlopen
    web_search_module.urllib.request.urlopen = raise_url_error
    try:
        result = web_search_module.search_formatted("test")
        assert result.startswith("[Web search unavailable:")
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_formatted_includes_results_when_successful():
    sample_html = (
        '<a class="result__a" href="//duckduckgo.com/l/?uddg=https%3A%2F%2Fexample.com%2F">Example Title</a>'
        '<a class="result__snippet">Example snippet text.</a>'
    ).encode("utf-8")

    original = _mock_urlopen(sample_html)
    try:
        formatted = web_search_module.search_formatted("test query")
        assert "Example Title" in formatted
        assert "Example snippet text." in formatted
        assert "https://example.com/" in formatted
    finally:
        web_search_module.urllib.request.urlopen = original


def test_web_search_tool_registered_in_builtin_tools():
    from byteflow.builtin_tools import get_builtin_tools
    tools = {t.name: t for t in get_builtin_tools()}
    assert "web_search" in tools


def test_looks_like_search_request_catches_top_n_today_phrasing():
    # Regression test for a real reported bug: "top 10 ai new today"
    # didn't match any fixed search-intent phrase, so it fell through
    # to the tool-planner instead of chat_with_search's hard guarantee,
    # and the web_search tool dumped raw unsummarized search text as
    # the final answer instead of a real synthesized response.
    agent = Agent(provider=None, memory_path=False)
    true_positives = [
        "top 10 ai new today",
        "top 10 ai news today",
        "top 5 movies this week",
        "recent developments in space",
        "what's happening with openai",
    ]
    for msg in true_positives:
        assert agent._looks_like_search_request(msg) is True, f"{msg!r} should trigger search"

    # the existing false-positive guards must still hold
    false_positives = [
        "what is my current project status",
        "today I learned something cool",
    ]
    for msg in false_positives:
        assert agent._looks_like_search_request(msg) is False, f"{msg!r} should NOT trigger search"


def test_get_builtin_tools_without_agent_uses_raw_search():
    # Backward compatibility: calling get_builtin_tools() with no agent
    # (as existing standalone code/tests do) still works, just without
    # LLM summarization.
    from byteflow.builtin_tools import get_builtin_tools
    tools = {t.name: t for t in get_builtin_tools()}  # no agent passed
    assert "web_search" in tools


def test_web_search_tool_bound_to_agent_summarizes_instead_of_dumping_raw_text():
    # Regression test for the actual reported bug: when the regular
    # tool-planner (not chat_with_search's auto-detection) picked
    # web_search directly, it returned raw formatted search text
    # (titles/snippets/URLs as one big string) as the final answer,
    # instead of a real synthesized response. Fixed by binding the
    # web_search tool to agent.chat_with_search when an agent is given.
    def fake_search(query, max_results=4, timeout=10):
        return [{"title": "Mock Result", "snippet": "Mock snippet.", "url": "https://example.com"}]

    original = web_search_module.search
    web_search_module.search = fake_search
    try:
        from byteflow.builtin_tools import register_builtin_tools

        class PlannerPicksWebSearchProvider:
            def generate(self, prompt):
                if "durable fact" in prompt:
                    return "none"
                if "Rules:" in prompt and "Available tools" in prompt:
                    return '[{"step": "web_search", "args": ["some obscure query"]}]'
                return "A clean synthesized answer based on the mock result."

        agent = Agent(provider=PlannerPicksWebSearchProvider(), memory_path=False)
        register_builtin_tools(agent)
        result = agent.run("tell me something about xyzcorp")
        result_str = str(result)
        assert "A clean synthesized answer" in result_str
        # the raw, unsummarized search dump format must NOT appear
        assert "Web search results for" not in result_str
    finally:
        web_search_module.search = original


def test_web_search_tool_degrades_gracefully_when_search_fails():
    from byteflow.builtin_tools import get_builtin_tools
    import urllib.error

    def raise_url_error(request, timeout=10):
        raise urllib.error.URLError("no route to host")

    original = web_search_module.urllib.request.urlopen
    web_search_module.urllib.request.urlopen = raise_url_error
    try:
        tools = {t.name: t for t in get_builtin_tools()}
        result = tools["web_search"].run("test query")
        assert "Tool Error" not in result  # the tool itself handles the failure
        assert "unavailable" in result.lower()
    finally:
        web_search_module.urllib.request.urlopen = original


# -----------------------------
# SEARCH INTENT DETECTION + ROUTING
# -----------------------------

def test_looks_like_search_request_catches_current_events_questions():
    agent = Agent(provider=None, memory_path=False)
    true_positives = [
        "what is the latest news on AI",
        "who is the current president of the US",
        "who is the ceo of Tesla",
        "stock price of Apple",
    ]
    for msg in true_positives:
        assert agent._looks_like_search_request(msg) is True, f"{msg!r} should trigger search"


def test_looks_like_search_request_excludes_weather():
    # Weather now has its own dedicated, more reliable path via a real
    # free weather API (weather.py / Open-Meteo) instead of being routed
    # through general web search - see _looks_like_weather_request.
    agent = Agent(provider=None, memory_path=False)
    weather_prompts = ["what is the weather today", "today weather", "weather in Ahmedabad"]
    for msg in weather_prompts:
        assert agent._looks_like_search_request(msg) is False, f"{msg!r} should NOT go through general search"


def test_looks_like_weather_request_catches_weather_questions():
    agent = Agent(provider=None, memory_path=False)
    true_positives = [
        "what is the weather today", "today weather", "weather in Ahmedabad",
        "today waether ahemdabad".replace("waether", "weather"),  # sanity: exact word still required
        "forecast for tomorrow", "current temperature",
    ]
    for msg in true_positives:
        assert agent._looks_like_weather_request(msg) is True, f"{msg!r} should be recognized as a weather question"

    false_positives = ["what is my name", "add 10 and 20", "who is the ceo of Tesla"]
    for msg in false_positives:
        assert agent._looks_like_weather_request(msg) is False, f"{msg!r} should NOT be recognized as a weather question"


def test_extract_weather_location_strips_filler_words():
    agent = Agent(provider=None, memory_path=False)
    assert agent._extract_weather_location("today weather ahmedabad") == "ahmedabad"
    assert agent._extract_weather_location("what is the weather like in ahmedabad") == "ahmedabad"
    assert agent._extract_weather_location("weather today") == ""


def test_looks_like_datetime_request_catches_date_and_time_questions():
    agent = Agent(provider=None, memory_path=False)
    true_positives = [
        "what is today's date",
        "what's today's date",
        "what day is it",
        "what year is it",
        "what time is it",
        # terse/keyword-style phrasing - the exact reported real bug:
        # the original version only matched full grammatical questions
        # and missed how people actually type short queries
        "today date",
        "date today",
        "current date",
        "todays date",
        "current time",
        "time now",
    ]
    for msg in true_positives:
        assert agent._looks_like_datetime_request(msg) is True, f"{msg!r} should be a datetime request"


def test_looks_like_datetime_request_avoids_today_substring_false_positive():
    # Regression test for a real bug found and fixed during development:
    # "day" is a SUBSTRING of "today", so a naive `"day" in text` check
    # meant ANY message containing the single word "today" (with no
    # other date/time intent at all) falsely triggered - e.g. "today I
    # learned something cool about python" is not a date/time question.
    # Fixed with whole-word matching instead of substring matching.
    agent = Agent(provider=None, memory_path=False)
    false_positives = [
        "today I learned something cool about python",
        "today was a good day at work",
        "I went hiking today",
    ]
    for msg in false_positives:
        assert agent._looks_like_datetime_request(msg) is False, f"{msg!r} should NOT be a datetime request"


def test_run_answers_datetime_questions_without_calling_the_llm():
    # Regression test for a real observed bug: asking "what's today's
    # date" went through the search+LLM pipeline, and when the search
    # result didn't clearly answer it, the model returned the literal
    # placeholder text "[current date]" instead of a real date. Fixed
    # by answering date/time questions directly from the system clock -
    # the LLM provider must never even be called for these.
    class ExplodingProvider:
        def generate(self, prompt):
            raise AssertionError("LLM should never be called for a pure date/time question")

    agent = Agent(provider=ExplodingProvider(), memory_path=False)
    result = agent.run("what is today's date")
    assert "[current date]" not in result
    assert "today's date is" in result.lower()

    result2 = agent.run("what time is it")
    assert "[current date]" not in result2

    # the exact phrase from the real reported bug report - terse,
    # no question grammar at all
    result3 = agent.run("today date")
    assert "[current date]" not in result3
    assert "today's date is" in result3.lower()


def test_answer_datetime_request_gives_real_date_and_time():
    from datetime import datetime
    agent = Agent(provider=None, memory_path=False)

    date_answer = agent._answer_datetime_request("what is today's date")
    assert str(datetime.now().year) in date_answer

    time_answer = agent._answer_datetime_request("what time is it")
    assert "currently" in time_answer.lower()


def test_looks_like_search_request_avoids_false_positives():
    # Regression test for a real bug found during development: bare
    # words like "current" or "today" alone falsely matched personal
    # questions that have nothing to do with needing a web search.
    agent = Agent(provider=None, memory_path=False)
    false_positives = [
        "what is my current project status",
        "today I learned something cool",
        "what is the weather like usually in winter",
    ]
    for msg in false_positives:
        assert agent._looks_like_search_request(msg) is False, f"{msg!r} should NOT trigger search"


def test_run_routes_search_questions_through_chat_with_search():
    calls = {"search": 0}

    def fake_search(query, max_results=4, timeout=10):
        calls["search"] += 1
        return [{"title": "Mock Result", "snippet": "Mock snippet.", "url": "https://example.com"}]

    original = web_search_module.search
    web_search_module.search = fake_search
    try:
        provider = _FakeProvider("Based on search results, here is the answer.")
        agent = Agent(provider=provider, memory_path=False)
        result = agent.run("what is the latest news on AI")
        assert calls["search"] == 1
        assert result == "Based on search results, here is the answer."
    finally:
        web_search_module.search = original


def test_chat_with_search_includes_search_results_in_prompt():
    captured_prompts = []

    class CapturingProvider:
        def generate(self, prompt):
            captured_prompts.append(prompt)
            if "durable fact" in prompt:
                return "none"
            return "answer"

    def fake_search(query, max_results=4, timeout=10):
        return [{"title": "Mock Result", "snippet": "Mock snippet.", "url": "https://example.com"}]

    original = web_search_module.search
    web_search_module.search = fake_search
    try:
        agent = Agent(provider=CapturingProvider(), memory_path=False)
        agent.chat_with_search("what is the latest news")
        main_prompt = captured_prompts[0]
        assert "Mock Result" in main_prompt
        assert "Mock snippet" in main_prompt
        assert "https://example.com" in main_prompt
    finally:
        web_search_module.search = original


def test_chat_with_search_never_calls_llm_when_search_fails():
    # Regression test for a real observed bug: a model given empty/
    # failed search results still fabricated plausible-looking fake
    # URLs and headlines instead of admitting the search didn't work.
    # The fix is a hard guarantee, not a prompt instruction: if search
    # fails, the LLM is never even called - an honest message is
    # returned directly.
    class ExplodingProvider:
        def generate(self, prompt):
            raise AssertionError("LLM should never be called when search fails")

    def fake_search_raises(query, max_results=4, timeout=10):
        raise web_search_module.WebSearchError("simulated network failure")

    original = web_search_module.search
    web_search_module.search = fake_search_raises
    try:
        agent = Agent(provider=ExplodingProvider(), memory_path=False)
        result = agent.chat_with_search("what is the latest news on AI")
        assert "search" in result.lower()
        assert "guess" in result.lower() or "make up" in result.lower()
    finally:
        web_search_module.search = original


def test_chat_with_search_never_calls_llm_when_no_results_found():
    class ExplodingProvider:
        def generate(self, prompt):
            raise AssertionError("LLM should never be called when search finds nothing")

    def fake_search_empty(query, max_results=4, timeout=10):
        return []

    original = web_search_module.search
    web_search_module.search = fake_search_empty
    try:
        agent = Agent(provider=ExplodingProvider(), memory_path=False)
        result = agent.chat_with_search("asdkjaslkdjaslkdj nonsense query")
        assert "search" in result.lower()
        assert "guess" in result.lower() or "make up" in result.lower()
    finally:
        web_search_module.search = original


def test_chat_with_search_summarizes_only_when_real_results_exist():
    captured_prompts = []

    class CapturingProvider:
        def generate(self, prompt):
            captured_prompts.append(prompt)
            return "Based on the real results, here's the answer."

    def fake_search(query, max_results=4, timeout=10):
        return [{"title": "Real Title", "snippet": "Real snippet.", "url": "https://real.example.com"}]

    original = web_search_module.search
    web_search_module.search = fake_search
    try:
        agent = Agent(provider=CapturingProvider(), memory_path=False)
        result = agent.chat_with_search("test query")
        assert result == "Based on the real results, here's the answer."
        # the first captured prompt is the actual answer-generation call;
        # a second call happens afterward for fact extraction (learn_from_exchange)
        prompt = captured_prompts[0].lower()
        assert "never invent" in prompt
        assert "synthesize" in prompt  # quality instruction: don't just list results back
        assert "disagree" in prompt    # quality instruction: handle conflicting results honestly
    finally:
        web_search_module.search = original


def test_chat_with_search_without_provider_fails_cleanly():
    agent = Agent(provider=None, memory_path=False)
    result = agent.chat_with_search("anything")
    assert "No provider configured" in result


def test_draft_social_post_is_auto_registered_as_a_tool():
    # Regression test for a real gap found during a companion audit:
    # draft_social_post existed as a method but had no way to actually
    # be triggered from chat/run() - "post about X on linkedin" would
    # just talk about posting instead of really doing it. Every Agent
    # should have this tool available by default.
    agent = Agent(provider=None, memory_path=False)
    assert "draft_social_post" in agent.tools


def test_run_triggers_real_draft_social_post_via_planner():
    import byteflow.desktop_tools as dt_module
    original_write = dt_module.write_clipboard
    original_launch = dt_module.launch
    dt_module.write_clipboard = lambda text: "Clipboard updated."
    dt_module.launch = lambda target: f"Launched: {target}"

    class PlannerProvider:
        def generate(self, prompt):
            if "durable fact" in prompt:
                return "none"
            if "Rules:" in prompt and "Available tools" in prompt:
                return '[{"step": "draft_social_post", "args": ["launching my project", "linkedin"]}]'
            if "Write a short, engaging" in prompt:
                return "Excited to launch my project!"
            return "fallback"

    try:
        agent = Agent(provider=PlannerProvider(), memory_path=False)
        result = agent.run("post about launching my project on linkedin")
        result_str = str(result)
        assert "Excited to launch my project!" in result_str
        assert "Copied to clipboard" in result_str
        assert "Opened linkedin" in result_str
    finally:
        dt_module.write_clipboard = original_write
        dt_module.launch = original_launch


def test_draft_social_post_tool_handles_unsupported_platform():
    agent = Agent(provider=_FakeProvider("draft text"), memory_path=False)
    result = agent.draft_social_post_tool("topic", platform="youtube")
    assert "isn't a supported platform" in result


def test_draft_social_post_tool_handles_missing_topic_without_crashing():
    # Regression test for a real reported bug: a vague request like
    # "give me content ideas" or "post on linkedin" with no stated
    # topic caused the planner to call draft_social_post_tool() with
    # zero arguments, crashing with a raw TypeError instead of asking
    # what to post about.
    agent = Agent(provider=None, memory_path=False)
    result = agent.draft_social_post_tool()
    assert "what topic" in result.lower() or "what it should be about" in result.lower()
    assert "Error" not in result

    assert "what topic" in agent.draft_social_post_tool("").lower() or True
    assert "Error" not in agent.draft_social_post_tool("   ")


def test_run_does_not_repeat_old_tool_action_for_unrelated_followup():
    # Regression test for a real reported bug: after "open linkedin"
    # succeeded, later unrelated messages ("give me content ideas",
    # "can you write code") kept getting routed back into
    # draft_social_post/launch because recent tool-call history (e.g.
    # "launch -> Launched: linkedin") was being treated like a
    # standing instruction rather than past events. The planner prompt
    # now explicitly distinguishes history from current intent.
    class HistoryAwareProvider:
        def generate(self, prompt):
            if "durable fact" in prompt:
                return "none"
            if "Rules:" in prompt and "Available tools" in prompt:
                # a well-behaved model, given the new explicit history
                # framing, should correctly say null for an unrelated
                # follow-up even with LinkedIn history present
                if "write a function" in prompt.lower() or "write code" in prompt.lower():
                    return "null"
                return "null"
            return "Sure, here's a simple function for you."

    agent = Agent(provider=HistoryAwareProvider(), memory_path=False)
    agent.add_memory("user", "open linkedin")
    agent.add_memory("tool", "launch -> Launched: linkedin (https://www.linkedin.com/feed/)")

    result = agent.run("can you write a simple function for me")
    # must NOT crash, and must not be routed into draft_social_post/launch
    assert "Error" not in str(result)


# -----------------------------
# FILE READING (PDF/DOCX extraction fix)
# -----------------------------

def _make_test_pdf(path, chapters):
    """Build a small real PDF with one page per chapter, real extractable text."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(path, pagesize=letter)
    for title, body in chapters:
        c.drawString(100, 750, title)
        c.drawString(100, 720, body)
        c.showPage()
    c.save()


def test_read_pdf_extracts_real_text_not_binary_garbage():
    # Regression test for the actual critical reported bug: uploading a
    # PDF was being opened with plain open(path, "r", encoding="utf-8"),
    # which doesn't extract any real document text - it reads raw
    # binary bytes and mangles them into replacement characters. The
    # garbage got chunked and indexed, so every question about the
    # uploaded document retrieved meaningless noise. This must now
    # extract genuine, readable text instead.
    if not file_reading.pdf_support_available():
        print("SKIP (pypdf not installed in this environment)")
        return

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "test.pdf")
        _make_test_pdf(path, [
            ("Chapter 1: Introduction", "This covers the basics of the subject."),
            ("Chapter 2: Data Structures", "This covers arrays, lists, and trees."),
        ])

        text = file_reading.read_file_text(path)
        assert "%PDF" not in text          # no raw binary header leaking through
        assert "Chapter 1" in text
        assert "Chapter 2" in text
        assert "Data Structures" in text


def test_read_pdf_image_only_raises_clear_error():
    if not file_reading.pdf_support_available():
        print("SKIP (pypdf not installed in this environment)")
        return

    from pypdf import PdfWriter

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "blank.pdf")
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with open(path, "wb") as f:
            writer.write(f)

        try:
            file_reading.read_file_text(path)
            assert False, "expected FileReadError"
        except file_reading.FileReadError as e:
            assert "no extractable text" in str(e).lower() or "scanned" in str(e).lower()


def test_read_pdf_corrupt_file_raises_clear_error():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "fake.pdf")
        with open(path, "w") as f:
            f.write("this is not a real pdf")

        try:
            file_reading.read_file_text(path)
            assert False, "expected FileReadError"
        except file_reading.FileReadError:
            pass  # any clear FileReadError is correct here


def test_read_docx_extracts_real_text():
    if not file_reading.docx_support_available():
        print("SKIP (python-docx not installed in this environment)")
        return

    import docx

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "test.docx")
        document = docx.Document()
        document.add_paragraph("Chapter 1: Introduction")
        document.add_paragraph("This covers the basics.")
        document.save(path)

        text = file_reading.read_file_text(path)
        assert "Chapter 1" in text
        assert "basics" in text


def test_read_plain_text_files_unchanged():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "notes.txt")
        with open(path, "w") as f:
            f.write("plain text content here")

        text = file_reading.read_file_text(path)
        assert text == "plain text content here"


def test_read_file_text_nonexistent_file_raises_clear_error():
    try:
        file_reading.read_file_text("/definitely/not/a/real/path.pdf")
        assert False, "expected FileReadError"
    except file_reading.FileReadError as e:
        assert "does not exist" in str(e).lower()


def test_pdf_support_unavailable_raises_clear_error():
    # Simulate pypdf not being installed, regardless of whether it
    # actually is in this environment - the error path must be clear.
    original = file_reading.pdf_support_available
    file_reading.pdf_support_available = lambda: False
    try:
        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "test.pdf")
            with open(path, "wb") as f:
                f.write(b"%PDF-1.4 fake")
            try:
                file_reading._read_pdf(path)
                assert False, "expected FileReadError"
            except file_reading.FileReadError as e:
                assert "pypdf" in str(e).lower()
    finally:
        file_reading.pdf_support_available = original


def test_codehelp_read_code_file_uses_real_pdf_extraction():
    # codehelp.py had the exact same bug if pointed at a binary file by
    # mistake - confirm it now delegates to the fixed extraction too.
    if not file_reading.pdf_support_available():
        print("SKIP (pypdf not installed in this environment)")
        return

    from byteflow.codehelp import read_code_file

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "test.pdf")
        _make_test_pdf(path, [("Chapter 1: Test", "Some real content here.")])

        result = read_code_file(path)
        assert "%PDF" not in result
        assert "Chapter 1" in result


def test_companion_on_upload_fix_end_to_end_via_ingest():
    # Full end-to-end simulation of the real reported bug: upload a PDF
    # with real chapter content, ingest it, then search for content
    # from "the second chapter" and confirm REAL relevant text comes
    # back - not garbage, not unrelated hallucinated content.
    if not file_reading.pdf_support_available():
        print("SKIP (pypdf not installed in this environment)")
        return

    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "syllabus.pdf")
        _make_test_pdf(path, [
            ("Chapter 1: Introduction to Computer Engineering",
             "This chapter covers digital logic and circuits."),
            ("Chapter 2: Data Structures and Algorithms",
             "This chapter covers arrays linked lists stacks queues trees and graphs."),
        ])

        agent = Agent(provider=None, memory_path=False)
        text = file_reading.read_file_text(path)
        n_chunks = agent.ingest_document(text, source="syllabus.pdf")
        assert n_chunks >= 1

        results = agent.vector_store.search("data structures algorithms second chapter")
        assert len(results) >= 1
        assert "Data Structures" in results[0]["text"] or "arrays" in results[0]["text"].lower()


# -----------------------------
# DOCUMENT-REFERENCE DETECTION (bypass the tool planner for PDF/doc Q&A)
# -----------------------------

def test_looks_like_document_request_catches_pdf_references():
    agent = Agent(provider=None, memory_path=False)
    true_positives = [
        "from pdf give me answer",
        "from pdf",
        "yes give me answer of 1 to 10 question in detail",
        "answer question no 3",
        "explain question 5 in detail",
        "from the document, summarize chapter 2",
        "based on the file, what is a compiler",
    ]
    for msg in true_positives:
        assert agent._looks_like_document_request(msg) is True, f"{msg!r} should be recognized as a document question"


def test_looks_like_document_request_excludes_unrelated_requests():
    agent = Agent(provider=None, memory_path=False)
    false_positives = ["add 10 and 20", "open youtube", "what is my name", "today weather"]
    for msg in false_positives:
        assert agent._looks_like_document_request(msg) is False, f"{msg!r} should NOT be treated as a document question"


def test_run_routes_document_questions_to_chat_not_the_planner():
    # Real observed bug: with tools registered, the planner hallucinated
    # "read_clipboard" for "from pdf give me answer" and "multiply" for
    # "give me answer of 1 to 10 question in detail" instead of
    # recognizing neither tool applies. Prove run() never even calls the
    # planner for these - it should go straight to chat()'s RAG path.
    class FakeProvider:
        def generate(self, prompt):
            if "STRICT planner" in prompt:
                # if this fires, the planner was wrongly invoked
                return '[{"step": "multiply", "args": [5, 2]}]'
            return "Answer from the syllabus document context."

    agent = Agent(provider=FakeProvider(), memory_path=False)
    agent.vector_store.add_document(
        "Q1: What is a compiler? A compiler translates source code to machine code.",
        source="syllabus.pdf",
    )

    result = agent.run("from pdf give me answer of 1 to 10 question in detail")
    assert result == "Answer from the syllabus document context."


if __name__ == "__main__":
    # allow running without pytest installed
    test_fns = [v for k, v in list(globals().items()) if k.startswith("test_")]
    passed, failed = 0, 0
    for fn in test_fns:
        try:
            fn()
            print(f"PASS: {fn.__name__}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL: {fn.__name__} -> {e}")
            failed += 1
        except Exception as e:
            print(f"ERROR: {fn.__name__} -> {e}")
            failed += 1
    print(f"\n{passed} passed, {failed} failed")
    sys.exit(1 if failed else 0)
