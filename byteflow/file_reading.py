"""
Real text extraction for uploaded files - fixes a critical bug where
binary document formats (PDF, DOCX) were being opened with plain
open(path, "r", encoding="utf-8", errors="replace"), which doesn't
extract any actual document text - it reads raw binary bytes (compressed
streams, font tables, internal structure) and mangles them into
replacement characters. The result got chunked and indexed into the
vector store as pure garbage, so every question about the uploaded
document retrieved meaningless noise instead of real content.

This module detects the actual file type and uses the right extraction
method:
  - .pdf  -> pypdf (optional dependency)
  - .docx -> python-docx (optional dependency)
  - everything else -> treated as plain text (the original behavior,
    which is correct for .txt, .py, .md, .csv, log files, etc.)

Both pypdf and python-docx are optional - if not installed, this
returns a clear, actionable error message instead of silently
producing garbage (the original failure mode) or crashing.
"""

import os


class FileReadError(Exception):
    pass


def pdf_support_available():
    try:
        import pypdf  # noqa: F401
        return True
    except ImportError:
        return False


def docx_support_available():
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _read_pdf(path):
    if not pdf_support_available():
        raise FileReadError(
            "Reading PDF files requires 'pypdf'. Install it with: pip install pypdf"
        )
    import pypdf

    try:
        reader = pypdf.PdfReader(path)
    except Exception as e:
        raise FileReadError(f"Could not open '{path}' as a PDF: {e}") from e

    pages_text = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            text = f"[Could not extract text from page {i + 1}: {e}]"
        pages_text.append(text)

    full_text = "\n\n".join(pages_text).strip()

    if not full_text:
        raise FileReadError(
            f"'{path}' appears to be a scanned/image-only PDF with no extractable "
            f"text (pypdf found {len(reader.pages)} page(s) but no text on any of "
            f"them). This needs OCR to read, which ByteFlow doesn't currently do."
        )

    return full_text


def _read_docx(path):
    if not docx_support_available():
        raise FileReadError(
            "Reading .docx files requires 'python-docx'. Install it with: pip install python-docx"
        )
    import docx

    try:
        document = docx.Document(path)
    except Exception as e:
        raise FileReadError(f"Could not open '{path}' as a .docx file: {e}") from e

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # also pull text out of tables, which paragraphs alone would miss
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    full_text = "\n".join(paragraphs).strip()

    if not full_text:
        raise FileReadError(f"'{path}' appears to be empty or contains no readable text.")

    return full_text


def _read_plain_text(path, max_chars=None):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    except OSError as e:
        raise FileReadError(f"Could not read '{path}': {e}") from e

    if max_chars and len(content) > max_chars:
        content = content[:max_chars] + "\n\n... [truncated] ..."

    return content


def read_file_text(path):
    """
    Extract the real text content of `path`, detecting its type by
    extension. Raises FileReadError with a clear message on any
    failure (missing file, unsupported/missing library, corrupt file,
    image-only PDF) - never silently returns binary garbage, which was
    the original bug.
    """
    if not os.path.isfile(path):
        raise FileReadError(f"'{path}' is not a file or does not exist.")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)

    # everything else (.txt, .py, .md, .csv, .json, log files, etc.)
    # is read as plain text, which is correct for genuinely text-based
    # formats - this is the original behavior, kept for non-binary files
    return _read_plain_text(path)
